import csv
import argparse
from pathlib import Path

from docx import Document
from openpyxl import Workbook


def write_csv(path: Path) -> None:
    rows = [
        {
            "Student ID": "A-001",
            "First Name": "aya",
            "Last Name": "BENALI",
            "Date of Birth": "2004-01-13",
            "Gender": "F",
            "Email": "Aya.Benali@Example.com",
            "Phone": "+213 555 000 111",
            "Institution": "Institute A",
            "Program": "Computer Science",
            "Enrollment Year": "2022",
        },
        {
            "Student ID": "A-002",
            "First Name": "youssef",
            "Last Name": " el  amrani ",
            "Date of Birth": "13/07/2003",
            "Gender": "Male",
            "Email": "youssef.amrani@example.com ",
            "Phone": "0555-000-222",
            "Institution": "Institute A",
            "Program": "Mathematics",
            "Enrollment Year": "2021",
        },
    ]

    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)


def write_excel(path: Path) -> None:
    # Use different column names on purpose to simulate heterogeneity
    rows = [
        {
            "matricule": "B-1001",
            "prenom": "Sara",
            "nom": "HADDAD",
            "dob": "2005/03/02",
            "sexe": "F",
            "mail": "sara.haddad@example.com",
            "tel": "(0555) 000-333",
            "etablissement": "Institute B",
            "filiere": "Physics",
            "annee_inscription": 2023,
        },
        {
            "matricule": "B-1002",
            "prenom": "Omar",
            "nom": "KHELIFI",
            "dob": "02-11-2004",
            "sexe": "M",
            "mail": "omar.khelifi@example.com",
            "tel": "0555000444",
            "etablissement": "Institute B",
            "filiere": "Chemistry",
            "annee_inscription": 2022,
        },
    ]

    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    headers = list(rows[0].keys())
    ws.append(headers)
    for row in rows:
        ws.append([row.get(h) for h in headers])
    wb.save(path)


def write_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("Student List - Institute C", level=1)

    table = doc.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "id"
    hdr_cells[1].text = "full_name"
    hdr_cells[2].text = "birthdate"
    hdr_cells[3].text = "email"
    hdr_cells[4].text = "gender"
    hdr_cells[5].text = "year"

    data = [
        ["C-77", "Nour Ait Ali", "2004.12.09", "nour.aitali@example.com", "Female", "2021"],
        ["C-78", "Mehdi  Boudiaf", "09/05/2003", "mehdi.boudiaf@example.com", "M", "2020"],
    ]

    for row in data:
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

    doc.add_paragraph("Institution: Institute C")
    doc.add_paragraph("Program: Civil Engineering")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate sample student files (CSV/XLSX/DOCX) for demo/testing.")
    parser.add_argument(
        "--output",
        default=None,
        help="Output directory (default: <repo>/sample_data)",
    )
    args = parser.parse_args()

    base = Path(__file__).resolve().parents[1]
    out_dir = Path(args.output).expanduser().resolve() if args.output else (base / "sample_data")

    write_csv(out_dir / "students_institute_a.csv")
    write_excel(out_dir / "students_institute_b.xlsx")
    write_docx(out_dir / "students_institute_c.docx")

    print(f"Wrote sample files to: {out_dir}")


if __name__ == "__main__":
    main()
