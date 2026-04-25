from __future__ import annotations

import argparse
import json
import re
import uuid
from dataclasses import dataclass
from datetime import datetime
from datetime import date
from pathlib import Path
from typing import Any, Iterable

from dateutil import parser as date_parser
from docx import Document
from openpyxl import load_workbook


SUPPORTED_EXTS = {".csv", ".xlsx", ".xls", ".docx"}


def _clean_str(value: Any) -> str | None:
    if value is None:
        return None
    # openpyxl can yield datetime/date/float/etc.
    if isinstance(value, (datetime, date)):
        return value.isoformat()
    s = str(value).strip()
    if s == "" or s.lower() in {"na", "n/a", "null", "none"}:
        return None
    return re.sub(r"\s+", " ", s)


def _normalize_gender(value: Any) -> str | None:
    s = _clean_str(value)
    if not s:
        return None
    s2 = s.strip().lower()
    if s2 in {"m", "male", "man", "masc", "masculin"}:
        return "M"
    if s2 in {"f", "female", "woman", "fem", "feminin", "féminin"}:
        return "F"
    return "Other"


def _normalize_email(value: Any) -> str | None:
    s = _clean_str(value)
    if not s:
        return None
    return s.strip().lower()


def _normalize_phone(value: Any) -> str | None:
    s = _clean_str(value)
    if not s:
        return None
    # keep digits and + only
    s = re.sub(r"[^0-9+]", "", s)
    return s or None


def _parse_date_to_iso(value: Any) -> str | None:
    s = _clean_str(value)
    if not s:
        return None
    try:
        dt = date_parser.parse(s, dayfirst=False, yearfirst=False).date()
    except Exception:
        return None
    # basic sanity: student DOB likely 1900..today
    today = date.today()
    if dt.year < 1900 or dt > today:
        return None
    return dt.isoformat()


def _title_name(value: Any) -> str | None:
    s = _clean_str(value)
    if not s:
        return None
    return s.title()


ALIAS_MAP: dict[str, str] = {
    # student id
    "student id": "student_id",
    "id": "student_id",
    "matricule": "student_id",
    "studentid": "student_id",
    "registration id": "student_id",
    # names
    "first name": "first_name",
    "firstname": "first_name",
    "prenom": "first_name",
    "prénom": "first_name",
    "given name": "first_name",
    "last name": "last_name",
    "lastname": "last_name",
    "nom": "last_name",
    "surname": "last_name",
    "family name": "last_name",
    "full_name": "full_name",
    "full name": "full_name",
    # dob
    "date of birth": "date_of_birth",
    "dob": "date_of_birth",
    "birthdate": "date_of_birth",
    "birth date": "date_of_birth",
    # contact
    "email": "email",
    "e-mail": "email",
    "mail": "email",
    "phone": "phone",
    "tel": "phone",
    "tél": "phone",
    "telephone": "phone",
    # misc
    "gender": "gender",
    "sexe": "gender",
    "institution": "institution",
    "etablissement": "institution",
    "établissement": "institution",
    "program": "program",
    "filiere": "program",
    "filière": "program",
    "enrollment year": "enrollment_year",
    # note: _normalize_columns converts underscores/hyphens to spaces
    "annee inscription": "enrollment_year",
    "année inscription": "enrollment_year",
    "year": "enrollment_year",
}


CANONICAL_FIELDS = [
    "student_id",
    "first_name",
    "last_name",
    "date_of_birth",
    "gender",
    "email",
    "phone",
    "institution",
    "program",
    "enrollment_year",
]


def _normalize_columns(columns: Iterable[str]) -> list[str]:
    out = []
    for c in columns:
        c0 = _clean_str(c) or ""
        c1 = re.sub(r"[_\-]+", " ", c0).strip().lower()
        c1 = re.sub(r"\s+", " ", c1)
        out.append(ALIAS_MAP.get(c1, c1))
    return out


def _split_full_name(full_name: str) -> tuple[str | None, str | None]:
    full_name = _clean_str(full_name) or ""
    if not full_name:
        return None, None
    parts = [p for p in full_name.split(" ") if p]
    if len(parts) == 1:
        return parts[0].title(), None
    first = " ".join(parts[:-1]).title()
    last = parts[-1].title()
    return first, last


def _generate_student_id(record: dict[str, Any]) -> str:
    seed = "|".join(
        [
            _clean_str(record.get("institution")) or "",
            _clean_str(record.get("first_name")) or "",
            _clean_str(record.get("last_name")) or "",
            _clean_str(record.get("date_of_birth")) or "",
            _clean_str(record.get("email")) or "",
        ]
    )
    return str(uuid.uuid5(uuid.NAMESPACE_URL, seed or str(uuid.uuid4())))


def _normalize_record(raw: dict[str, Any], default_institution: str | None) -> dict[str, Any]:
    record: dict[str, Any] = {}

    student_id = _clean_str(raw.get("student_id"))
    full_name = _clean_str(raw.get("full_name"))

    first_name = _title_name(raw.get("first_name"))
    last_name = _title_name(raw.get("last_name"))

    if (not first_name or not last_name) and full_name:
        f, l = _split_full_name(full_name)
        first_name = first_name or f
        last_name = last_name or l

    institution = _clean_str(raw.get("institution")) or (default_institution.strip() if default_institution else None)

    record["student_id"] = student_id
    record["first_name"] = first_name
    record["last_name"] = last_name
    record["date_of_birth"] = _parse_date_to_iso(raw.get("date_of_birth"))
    record["gender"] = _normalize_gender(raw.get("gender"))
    record["email"] = _normalize_email(raw.get("email"))
    record["phone"] = _normalize_phone(raw.get("phone"))
    record["institution"] = institution
    record["program"] = _clean_str(raw.get("program"))

    enrollment_year = _clean_str(raw.get("enrollment_year"))
    if enrollment_year and enrollment_year.isdigit():
        record["enrollment_year"] = int(enrollment_year)
    else:
        record["enrollment_year"] = None

    if not record["student_id"]:
        record["student_id"] = _generate_student_id(record)

    # remove keys with None to keep JSON compact
    return {k: v for k, v in record.items() if v is not None}


@dataclass(frozen=True)
class ExtractResult:
    rows: list[dict[str, Any]]
    warnings: list[str]


def _extract_from_csv(path: Path) -> ExtractResult:
    import csv

    with path.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        if not reader.fieldnames:
            return ExtractResult(rows=[], warnings=["CSV has no header row"])
        fieldnames = _normalize_columns(reader.fieldnames)
        rows: list[dict[str, Any]] = []
        for raw in reader:
            # Map raw header → normalized header
            mapped: dict[str, Any] = {}
            for i, original_key in enumerate(reader.fieldnames):
                mapped[fieldnames[i]] = raw.get(original_key)
            rows.append(mapped)
    return ExtractResult(rows=rows, warnings=[])


def _extract_from_excel(path: Path) -> ExtractResult:
    wb = load_workbook(filename=path, read_only=True, data_only=True)
    try:
        ws = wb.active

        values = list(ws.iter_rows(values_only=True))
        if not values:
            return ExtractResult(rows=[], warnings=["Excel sheet is empty"])

        header_row = values[0]
        headers = _normalize_columns([str(h) if h is not None else "" for h in header_row])

        rows: list[dict[str, Any]] = []
        for row in values[1:]:
            if row is None:
                continue
            mapped = {headers[i] if i < len(headers) else f"col_{i}": row[i] for i in range(len(row))}
            # drop rows that are fully empty
            if all(_clean_str(v) is None for v in mapped.values()):
                continue
            rows.append(mapped)

        return ExtractResult(rows=rows, warnings=[])
    finally:
        # important on Windows so temporary XLSX files can be deleted
        wb.close()


def _extract_from_docx(path: Path) -> ExtractResult:
    doc = Document(path)
    warnings: list[str] = []

    # 1) tables → rows
    rows: list[dict[str, Any]] = []
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        headers = _normalize_columns(headers)

        for r in table.rows[1:]:
            values = [cell.text.strip() for cell in r.cells]
            row = {headers[i] if i < len(headers) else f"col_{i}": values[i] for i in range(len(values))}
            rows.append(row)

    # 2) try to infer institution/program from paragraphs (optional)
    inferred: dict[str, str] = {}
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text or ":" not in text:
            continue
        k, v = text.split(":", 1)
        k_norm = (_normalize_columns([k])[0] or "").strip().lower()
        v = v.strip()
        if k_norm in {"institution", "program"} and v:
            inferred[k_norm] = v

    if inferred and rows:
        for row in rows:
            row.setdefault("institution", inferred.get("institution"))
            row.setdefault("program", inferred.get("program"))

    if not rows:
        warnings.append("No tables found in DOCX; nothing extracted")

    return ExtractResult(rows=rows, warnings=warnings)


def extract_rows(path: Path) -> ExtractResult:
    ext = path.suffix.lower()
    if ext == ".csv":
        return _extract_from_csv(path)
    if ext in {".xlsx", ".xls"}:
        return _extract_from_excel(path)
    if ext == ".docx":
        return _extract_from_docx(path)
    raise ValueError(f"Unsupported file extension: {ext}")


def iter_input_files(input_path: Path) -> list[Path]:
    if input_path.is_file():
        return [input_path]
    files: list[Path] = []
    for p in input_path.rglob("*"):
        if p.is_file() and p.suffix.lower() in SUPPORTED_EXTS:
            files.append(p)
    return sorted(files)


def normalize_file(input_file: Path, out_dir: Path, default_institution: str | None) -> dict[str, dict[str, Any]]:
    extracted = extract_rows(input_file)

    normalized: dict[str, dict[str, Any]] = {}
    for raw in extracted.rows:
        # normalize keys too (docx table may have already-normalized headers)
        raw2 = {k: v for k, v in raw.items()}
        # If there are non-canonical keys, still allow them, but only use known ones
        picked = {k: raw2.get(k) for k in set(CANONICAL_FIELDS + ["full_name"]) if k in raw2}
        record = _normalize_record(picked, default_institution)
        sid = record.get("student_id")
        if not sid:
            # should not happen due to generation, but keep safe
            sid = str(uuid.uuid4())
            record["student_id"] = sid
        normalized[sid] = record

    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{input_file.stem}.json"
    payload = {
        "source_file": str(input_file),
        "students": normalized,
        "warnings": extracted.warnings,
    }
    out_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return normalized


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Normalize student records from CSV/Excel/Word (.docx) into JSON."
    )
    parser.add_argument("--input", required=True, help="Input file or folder containing CSV/XLSX/DOCX")
    parser.add_argument("--output", required=True, help="Output folder for normalized JSON")
    parser.add_argument(
        "--default-institution",
        default=None,
        help="Fallback institution name used when source data has no institution column",
    )
    args = parser.parse_args()

    input_path = Path(args.input).expanduser().resolve()
    out_dir = Path(args.output).expanduser().resolve()

    files = iter_input_files(input_path)
    if not files:
        raise SystemExit(f"No supported files found under: {input_path}")

    out_dir.mkdir(parents=True, exist_ok=True)

    total_students = 0
    for f in files:
        normalized = normalize_file(f, out_dir=out_dir, default_institution=args.default_institution)
        total_students += len(normalized)

    print(f"Processed {len(files)} file(s); normalized {total_students} student record(s)")
    print(f"Output written to: {out_dir}")


if __name__ == "__main__":
    main()
